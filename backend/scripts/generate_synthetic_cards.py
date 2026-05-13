"""
generate_synthetic_cards.py — Synthetic МЧС РК operational fire safety card generator.

Usage:
    python backend/scripts/generate_synthetic_cards.py
    python backend/scripts/generate_synthetic_cards.py --count 35 --output-dir backend/data/sample/synthetic_cards

Generates 35 realistic-looking synthetic cards:
    - 15 text PDFs (reportlab)
    - 10 scan-like JPGs (PDF rendered to PIL image + gaussian noise + slight rotation)
    - 5 DOCX files (python-docx)
    - 5 intentionally incomplete/bad quality cards (edge cases)

All user-facing text is in Russian.
"""

import argparse
import math
import os
import random
import shutil
import sys
from datetime import date, timedelta
from pathlib import Path

# ---------------------------------------------------------------------------
# Data pools (all Russian)
# ---------------------------------------------------------------------------

BUILDING_NAMES = [
    ('ТРЦ "Хан Шатыр"', 'Ф2.1'),
    ('ТРЦ "Мега Астана"', 'Ф2.1'),
    ('Школа №45', 'Ф4.1'),
    ('Школа №112', 'Ф4.1'),
    ('Школа №78', 'Ф4.1'),
    ('Детский сад №12', 'Ф1.1'),
    ('Детский сад №34', 'Ф1.1'),
    ('ЖК "Астана"', 'Ф1.3'),
    ('ЖК "Нур-Сити"', 'Ф1.3'),
    ('ЖК "Триумф"', 'Ф1.3'),
    ('Завод ТОО "СтройМаш"', 'Ф5.1'),
    ('Завод ТОО "АстанаМеталл"', 'Ф5.1'),
    ('Завод ТОО "КазПром"', 'Ф5.2'),
    ('Больница №3', 'Ф1.1'),
    ('Больница им. Бурденко', 'Ф1.1'),
    ('Офисный центр "Байтерек"', 'Ф4.2'),
    ('Офисный центр "Думан"', 'Ф4.2'),
    ('Бизнес-центр "Ак Орда Плаза"', 'Ф4.2'),
    ('Гостиница "Казахстан Хилтон"', 'Ф1.2'),
    ('Гостиница "Рамада"', 'Ф1.2'),
    ('Склад ТОО "АстанаЛогист"', 'Ф5.2'),
    ('Кинотеатр "Синема Сити"', 'Ф2.2'),
    ('Дворец спорта "Барыс"', 'Ф2.1'),
    ('Библиотека НАН РК', 'Ф2.1'),
    ('Колледж технологий №5', 'Ф4.2'),
]

STREETS = [
    'пр. Туран', 'ул. Сейткали Мендешев', 'пр. Достык',
    'ул. Бейбітшілік', 'пр. Тәуелсіздік', 'ул. Кунаева',
    'ул. Иманова', 'пр. Республики', 'ул. Аблай хана',
    'ул. Сарыарқа', 'пр. Нұр-Жол', 'ул. Бейбарыс',
]

DISTRICTS = ['Есіл', 'Алматы', 'Байқоңыр', 'Сарыарқа', 'Нұра']

WALL_MATERIALS = [
    'кирпич', 'монолит', 'металлокаркас', 'дерево', 'панельный'
]

FIRE_RESISTANCE = ['I', 'II', 'III', 'IV', 'V']

ALARM_SYSTEMS = [
    'Bosch FPA-5000', 'Болид', 'Орион', 'Siemens Cerberus',
]

RESPONSIBLE_NAMES = [
    ('Ахметов Серік Болатұлы', '+7 701 234 56 78'),
    ('Нұрланов Дәурен Асқарович', '+7 702 345 67 89'),
    ('Сейткали Алия Маратовна', '+7 707 456 78 90'),
    ('Бекова Гүлнар Ерланқызы', '+7 705 567 89 01'),
    ('Жаксыбеков Рустем Нурланович', '+7 771 678 90 12'),
    ('Асанова Карлыгаш Олжасовна', '+7 777 789 01 23'),
    ('Мухамедов Азамат Серикович', '+7 747 890 12 34'),
    ('Турсунова Динара Бекжановна', '+7 700 901 23 45'),
]

HAZARDS_GAS = ['природный газ', 'СУГ (сжиженный углеводородный газ)']
HAZARDS_ELECTRICAL = ['0,4 кВ', '6 кВ', '10 кВ']

ROAD_DESCRIPTIONS = [
    'Асфальтированная дорога шириной 6 м, подъезд с северной стороны',
    'Грунтовая дорога, ограниченный проезд в зимнее время',
    'Асфальтированная дорога шириной 8 м, подъезд с западной стороны',
    'Бетонная площадка, круговой объезд',
    'Асфальтированная дорога шириной 7 м, одностороннее движение',
    'Подъездной путь со стороны главного входа, ширина 5 м',
]

WATER_SOURCES = [
    'Пожарный водоём (V=200 м³, расстояние 150 м)',
    'Река Есіл, расстояние 400 м',
    'Внутренний противопожарный водопровод ∅100 мм',
    'Пожарный резервуар V=500 м³',
]

# ---------------------------------------------------------------------------
# Helper generators
# ---------------------------------------------------------------------------

def rnd_date(year_from: int, year_to: int) -> date:
    start = date(year_from, 1, 1)
    end = date(year_to, 12, 31)
    delta = (end - start).days
    return start + timedelta(days=random.randint(0, delta))


def card_number(idx: int) -> str:
    year = random.randint(2019, 2024)
    station = random.randint(1, 20)
    return f"{year}-ПЧ{station}-{idx:03d}"


def generate_card_data(idx: int, bad: bool = False) -> dict:
    """Generate a dict of card fields. If bad=True, omit 30-50% of fields."""
    name, category = random.choice(BUILDING_NAMES)
    street = random.choice(STREETS)
    house_num = random.randint(1, 200)
    district = random.choice(DISTRICTS)
    responsible_name, responsible_phone = random.choice(RESPONSIBLE_NAMES)

    approval_date = rnd_date(2019, 2024)
    # Some revisions are intentionally outdated (>3 years ago)
    if random.random() < 0.4 or bad:
        revision_date = rnd_date(2018, 2021)
    else:
        revision_date = rnd_date(2022, 2024)

    floors = random.randint(1, 25)
    area = random.randint(500, 150000)
    wall_material = random.choice(WALL_MATERIALS)
    year_built = random.randint(1960, 2023)
    fire_resistance = random.choice(FIRE_RESISTANCE)

    has_alarm = random.random() > 0.1
    alarm_type = random.choice(ALARM_SYSTEMS) if has_alarm else None
    has_sprinkler = random.random() > 0.4
    has_smoke_removal = random.random() > 0.3
    evacuation_exits = random.randint(2, 8)

    hydrant_distance = random.randint(20, 200)
    hydrant_status = random.choice(['рабочий', 'нерабочий'])

    has_gas = random.random() > 0.5
    gas_type = random.choice(HAZARDS_GAS) if has_gas else None
    has_lvg = random.random() > 0.6
    lvg_volume = round(random.uniform(0.5, 50.0), 1) if has_lvg else None
    has_electrical = random.random() > 0.4
    electrical_class = random.choice(HAZARDS_ELECTRICAL) if has_electrical else None

    roads = [random.choice(ROAD_DESCRIPTIONS) for _ in range(random.randint(1, 3))]

    extra_water = random.choice(WATER_SOURCES) if random.random() > 0.5 else None

    card = {
        'card_number': card_number(idx),
        'approval_date': approval_date.strftime('%d.%m.%Y'),
        'revision_date': revision_date.strftime('%d.%m.%Y'),
        'object_name': name,
        'address': f'{street}, {house_num}',
        'district': district,
        'category': category,
        'responsible_name': responsible_name,
        'responsible_phone': responsible_phone,
        'floors': floors,
        'area_m2': area,
        'wall_material': wall_material,
        'year_built': year_built,
        'fire_resistance': fire_resistance,
        'has_alarm': has_alarm,
        'alarm_type': alarm_type,
        'has_sprinkler': has_sprinkler,
        'has_smoke_removal': has_smoke_removal,
        'evacuation_exits': evacuation_exits,
        'hydrant_distance_m': hydrant_distance,
        'hydrant_status': hydrant_status,
        'extra_water': extra_water,
        'has_gas': has_gas,
        'gas_type': gas_type,
        'has_lvg': has_lvg,
        'lvg_volume_m3': lvg_volume,
        'has_electrical': has_electrical,
        'electrical_class': electrical_class,
        'roads': roads,
    }

    if bad:
        # Drop 30-50% of optional fields randomly
        optional_keys = [
            'revision_date', 'category', 'responsible_phone',
            'floors', 'area_m2', 'wall_material', 'year_built',
            'fire_resistance', 'alarm_type', 'has_sprinkler',
            'has_smoke_removal', 'evacuation_exits',
            'hydrant_distance_m', 'hydrant_status',
            'extra_water', 'has_gas', 'gas_type',
            'has_lvg', 'lvg_volume_m3', 'has_electrical',
            'electrical_class', 'roads',
        ]
        n_drop = random.randint(
            max(1, len(optional_keys) // 3),
            max(2, len(optional_keys) // 2)
        )
        for key in random.sample(optional_keys, n_drop):
            card[key] = None

    return card


def card_lines(card: dict) -> list[str]:
    """Return a list of human-readable text lines for the card."""
    lines = []

    def add(label: str, value) -> None:
        if value is not None:
            lines.append(f'{label}: {value}')
        else:
            lines.append(f'{label}: [не указано]')

    lines.append('=' * 60)
    lines.append('ОПЕРАТИВНАЯ КАРТОЧКА ПОЖАРНОЙ БЕЗОПАСНОСТИ')
    lines.append('Министерство по чрезвычайным ситуациям Республики Казахстан')
    lines.append('=' * 60)
    lines.append('')
    add('Номер карточки', card.get('card_number'))
    add('Дата утверждения', card.get('approval_date'))
    add('Дата последней ревизии', card.get('revision_date'))
    lines.append('')
    lines.append('РАЗДЕЛ 1. СВЕДЕНИЯ ОБ ОБЪЕКТЕ')
    lines.append('-' * 40)
    add('Название объекта', card.get('object_name'))
    add('Адрес', card.get('address'))
    add('Район', card.get('district'))
    add('Функциональная категория', card.get('category'))
    add('Ответственное лицо', card.get('responsible_name'))
    add('Телефон ответственного', card.get('responsible_phone'))
    lines.append('')
    lines.append('РАЗДЕЛ 2. ХАРАКТЕРИСТИКИ ЗДАНИЯ')
    lines.append('-' * 40)
    add('Количество этажей', card.get('floors'))
    add('Общая площадь (м²)', card.get('area_m2'))
    add('Материал стен', card.get('wall_material'))
    add('Год постройки', card.get('year_built'))
    add('Степень огнестойкости', card.get('fire_resistance'))
    lines.append('')
    lines.append('РАЗДЕЛ 3. СИСТЕМЫ ПОЖАРНОЙ БЕЗОПАСНОСТИ')
    lines.append('-' * 40)
    alarm_val = card.get('alarm_type')
    has_alarm = card.get('has_alarm')
    if has_alarm is None:
        lines.append('Пожарная сигнализация: [не указано]')
    elif has_alarm:
        lines.append(f'Пожарная сигнализация: есть (тип: {alarm_val or "не указан"})')
    else:
        lines.append('Пожарная сигнализация: нет')

    sprinkler = card.get('has_sprinkler')
    if sprinkler is None:
        lines.append('АУПТ (спринклер): [не указано]')
    else:
        lines.append(f'АУПТ (спринклер): {"да" if sprinkler else "нет"}')

    smoke = card.get('has_smoke_removal')
    if smoke is None:
        lines.append('Система дымоудаления: [не указано]')
    else:
        lines.append(f'Система дымоудаления: {"да" if smoke else "нет"}')

    add('Количество эвакуационных выходов', card.get('evacuation_exits'))
    lines.append('')
    lines.append('РАЗДЕЛ 4. ВОДОСНАБЖЕНИЕ')
    lines.append('-' * 40)
    h_dist = card.get('hydrant_distance_m')
    h_status = card.get('hydrant_status')
    if h_dist is not None and h_status is not None:
        lines.append(f'Ближайший гидрант: расстояние {h_dist} м, статус — {h_status}')
    else:
        lines.append('Ближайший гидрант: [не указано]')
    extra = card.get('extra_water')
    if extra:
        lines.append(f'Дополнительный источник воды: {extra}')
    lines.append('')
    lines.append('РАЗДЕЛ 5. ОПАСНЫЕ ФАКТОРЫ')
    lines.append('-' * 40)
    has_gas = card.get('has_gas')
    gas_type = card.get('gas_type')
    if has_gas:
        lines.append(f'Газовые системы: {gas_type or "[тип не указан]"}')
    elif has_gas is False:
        lines.append('Газовые системы: отсутствуют')
    else:
        lines.append('Газовые системы: [не указано]')

    has_lvg = card.get('has_lvg')
    lvg_vol = card.get('lvg_volume_m3')
    if has_lvg:
        lines.append(f'ЛВЖ: присутствуют (объём: {lvg_vol or "[не указан]"} м³)')
    elif has_lvg is False:
        lines.append('ЛВЖ: отсутствуют')
    else:
        lines.append('ЛВЖ: [не указано]')

    has_el = card.get('has_electrical')
    el_class = card.get('electrical_class')
    if has_el:
        lines.append(f'Электроустановки: {el_class or "[класс не указан]"}')
    elif has_el is False:
        lines.append('Электроустановки: нет сведений')
    else:
        lines.append('Электроустановки: [не указано]')

    lines.append('')
    lines.append('РАЗДЕЛ 6. ПОДЪЕЗДНЫЕ ПУТИ')
    lines.append('-' * 40)
    roads = card.get('roads')
    if roads:
        for i, road in enumerate(roads, 1):
            lines.append(f'Дорога {i}: {road}')
    else:
        lines.append('[Данные о подъездных путях отсутствуют]')

    lines.append('')
    lines.append('=' * 60)
    lines.append('Карточка составлена в соответствии с приказом МЧС РК №242')
    lines.append('=' * 60)
    return lines


# ---------------------------------------------------------------------------
# PDF generator (reportlab)
# ---------------------------------------------------------------------------

def generate_pdf(card: dict, output_path: str) -> None:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.utils import simpleSplit

    # Use a built-in font that supports basic Latin + Cyrillic via encoding trick.
    # reportlab's built-in fonts do NOT include Cyrillic. We must embed a TTF.
    # Try to find a system Cyrillic font; fall back to drawing bytes directly.
    font_name = 'Helvetica'
    font_bold = 'Helvetica-Bold'

    # Try to register a Cyrillic TTF if available on macOS / Linux
    cyrillic_candidates = [
        '/System/Library/Fonts/Supplemental/Arial.ttf',        # macOS
        '/Library/Fonts/Arial.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',      # Linux
        '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
        '/usr/share/fonts/truetype/freefont/FreeSans.ttf',
    ]
    cyrillic_bold_candidates = [
        '/System/Library/Fonts/Supplemental/Arial Bold.ttf',
        '/Library/Fonts/Arial Bold.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
    ]

    for path in cyrillic_candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont('CyrillicFont', path))
                font_name = 'CyrillicFont'
            except Exception:
                pass
            break

    for path in cyrillic_bold_candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont('CyrillicFontBold', path))
                font_bold = 'CyrillicFontBold'
            except Exception:
                pass
            break

    width, height = A4
    c = canvas.Canvas(output_path, pagesize=A4)
    margin = 40
    y = height - margin
    line_height = 14

    lines = card_lines(card)
    for line in lines:
        if y < margin + line_height:
            c.showPage()
            y = height - margin

        if line.startswith('===') or line.startswith('РАЗДЕЛ') or line.startswith('ОПЕРАТИВНАЯ') or line.startswith('Министерство'):
            c.setFont(font_bold, 10)
        elif line.startswith('---'):
            c.setFont(font_name, 8)
        else:
            c.setFont(font_name, 9)

        # Word-wrap long lines
        max_chars = int((width - 2 * margin) / 5.5)
        wrapped = simpleSplit(line, font_name, 9, width - 2 * margin)
        for wrapped_line in wrapped:
            if y < margin + line_height:
                c.showPage()
                y = height - margin
            c.drawString(margin, y, wrapped_line)
            y -= line_height

    c.save()


# ---------------------------------------------------------------------------
# DOCX generator (python-docx)
# ---------------------------------------------------------------------------

def generate_docx(card: dict, output_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading('ОПЕРАТИВНАЯ КАРТОЧКА ПОЖАРНОЙ БЕЗОПАСНОСТИ', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        'Министерство по чрезвычайным ситуациям Республики Казахстан'
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    def add_field(label: str, value) -> None:
        p = doc.add_paragraph()
        run_label = p.add_run(f'{label}: ')
        run_label.bold = True
        if value is not None:
            p.add_run(str(value))
        else:
            run_missing = p.add_run('[не указано]')
            run_missing.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    def add_section(title_text: str) -> None:
        doc.add_paragraph()
        heading = doc.add_heading(title_text, level=2)

    add_field('Номер карточки', card.get('card_number'))
    add_field('Дата утверждения', card.get('approval_date'))
    add_field('Дата последней ревизии', card.get('revision_date'))

    add_section('1. Сведения об объекте')
    add_field('Название объекта', card.get('object_name'))
    add_field('Адрес', card.get('address'))
    add_field('Район', card.get('district'))
    add_field('Функциональная категория', card.get('category'))
    add_field('Ответственное лицо', card.get('responsible_name'))
    add_field('Телефон ответственного', card.get('responsible_phone'))

    add_section('2. Характеристики здания')
    add_field('Количество этажей', card.get('floors'))
    add_field('Общая площадь (м²)', card.get('area_m2'))
    add_field('Материал стен', card.get('wall_material'))
    add_field('Год постройки', card.get('year_built'))
    add_field('Степень огнестойкости', card.get('fire_resistance'))

    add_section('3. Системы пожарной безопасности')
    has_alarm = card.get('has_alarm')
    alarm_type = card.get('alarm_type')
    if has_alarm is None:
        add_field('Пожарная сигнализация', None)
    elif has_alarm:
        add_field('Пожарная сигнализация', f'есть (тип: {alarm_type or "не указан"})')
    else:
        add_field('Пожарная сигнализация', 'нет')

    sprinkler = card.get('has_sprinkler')
    add_field('АУПТ (спринклер)', 'да' if sprinkler else ('нет' if sprinkler is False else None))
    smoke = card.get('has_smoke_removal')
    add_field('Система дымоудаления', 'да' if smoke else ('нет' if smoke is False else None))
    add_field('Количество эвакуационных выходов', card.get('evacuation_exits'))

    add_section('4. Водоснабжение')
    h_dist = card.get('hydrant_distance_m')
    h_status = card.get('hydrant_status')
    if h_dist is not None and h_status is not None:
        add_field('Ближайший гидрант', f'{h_dist} м, статус — {h_status}')
    else:
        add_field('Ближайший гидрант', None)
    add_field('Дополнительный источник воды', card.get('extra_water'))

    add_section('5. Опасные факторы')
    has_gas = card.get('has_gas')
    gas_type = card.get('gas_type')
    if has_gas:
        add_field('Газовые системы', gas_type or '[тип не указан]')
    elif has_gas is False:
        add_field('Газовые системы', 'отсутствуют')
    else:
        add_field('Газовые системы', None)

    has_lvg = card.get('has_lvg')
    lvg_vol = card.get('lvg_volume_m3')
    if has_lvg:
        add_field('ЛВЖ', f'присутствуют (объём: {lvg_vol} м³)')
    elif has_lvg is False:
        add_field('ЛВЖ', 'отсутствуют')
    else:
        add_field('ЛВЖ', None)

    has_el = card.get('has_electrical')
    el_class = card.get('electrical_class')
    if has_el:
        add_field('Электроустановки', el_class or '[класс не указан]')
    elif has_el is False:
        add_field('Электроустановки', 'нет сведений')
    else:
        add_field('Электроустановки', None)

    add_section('6. Подъездные пути')
    roads = card.get('roads')
    if roads:
        for i, road in enumerate(roads, 1):
            add_field(f'Дорога {i}', road)
    else:
        doc.add_paragraph('[Данные о подъездных путях отсутствуют]')

    doc.add_paragraph()
    footer = doc.add_paragraph('Карточка составлена в соответствии с приказом МЧС РК №242')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)


# ---------------------------------------------------------------------------
# Scan-like JPG generator (PIL — no pdf2image / poppler needed)
# ---------------------------------------------------------------------------

def generate_scan_jpg(card: dict, output_path: str) -> None:
    """
    Render card text onto a white A4-sized PIL image, then add:
      - Gaussian noise (σ ≈ 15)
      - Slight rotation (±3°)
    Saves as JPEG.
    No poppler / pdf2image dependency needed.
    """
    from PIL import Image, ImageDraw, ImageFont, ImageFilter
    import numpy as np

    # A4 at 150 DPI
    width, height = 1240, 1754
    img = Image.new('RGB', (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    # Try to get a TTF font; fall back to default bitmap font
    font_regular = None
    font_bold = None
    font_size_regular = 18
    font_size_bold = 20

    font_candidates = [
        '/System/Library/Fonts/Supplemental/Arial.ttf',
        '/Library/Fonts/Arial.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
        '/usr/share/fonts/truetype/freefont/FreeSans.ttf',
    ]
    bold_candidates = [
        '/System/Library/Fonts/Supplemental/Arial Bold.ttf',
        '/Library/Fonts/Arial Bold.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
    ]

    for path in font_candidates:
        if os.path.exists(path):
            try:
                from PIL import ImageFont
                font_regular = ImageFont.truetype(path, font_size_regular)
                break
            except Exception:
                pass

    for path in bold_candidates:
        if os.path.exists(path):
            try:
                from PIL import ImageFont
                font_bold = ImageFont.truetype(path, font_size_bold)
                break
            except Exception:
                pass

    if font_regular is None:
        font_regular = ImageFont.load_default()
    if font_bold is None:
        font_bold = font_regular

    margin = 60
    x = margin
    y = margin
    line_h = font_size_regular + 6

    lines = card_lines(card)
    for line in lines:
        if y > height - margin:
            break

        is_heading = (
            line.startswith('===') or
            line.startswith('РАЗДЕЛ') or
            line.startswith('ОПЕРАТИВНАЯ') or
            line.startswith('Министерство')
        )
        font = font_bold if is_heading else font_regular
        color = (0, 0, 0)

        # Simple manual wrap
        max_width = width - 2 * margin
        words = line.split(' ')
        current_line = ''
        for word in words:
            test_line = (current_line + ' ' + word).strip()
            # Estimate width: ~10px per char at font_size 18
            if len(test_line) * (font_size_regular * 0.55) > max_width and current_line:
                draw.text((x, y), current_line, fill=color, font=font)
                y += line_h
                current_line = word
            else:
                current_line = test_line
        if current_line:
            draw.text((x, y), current_line, fill=color, font=font)
        y += line_h

    # Add gaussian noise
    img_array = np.array(img, dtype=np.float32)
    noise = np.random.normal(0, 15, img_array.shape)
    img_array = np.clip(img_array + noise, 0, 255).astype(np.uint8)
    img = Image.fromarray(img_array)

    # Add slight rotation (-3° to +3°)
    angle = random.uniform(-3.0, 3.0)
    img = img.rotate(angle, expand=False, fillcolor=(255, 255, 255))

    # Slight blur to simulate scan
    img = img.filter(ImageFilter.GaussianBlur(radius=0.5))

    img.save(output_path, 'JPEG', quality=85)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description='Generate synthetic МЧС РК operational fire safety cards.'
    )
    parser.add_argument(
        '--count', type=int, default=35,
        help='Total number of cards to generate (default: 35)'
    )
    parser.add_argument(
        '--output-dir', type=str,
        default='backend/data/sample/synthetic_cards',
        help='Output directory (default: backend/data/sample/synthetic_cards)'
    )
    args = parser.parse_args()

    total = args.count
    output_dir = Path(args.output_dir)

    # Plan: 15 PDFs, 10 scan JPGs, 5 DOCXs, 5 bad cards
    # bad cards are distributed: 3 bad PDFs + 2 bad DOCX (or similar)
    n_pdf = 15
    n_jpg = 10
    n_docx = 5
    n_bad = 5

    # Adjust if total != 35
    if total != 35:
        ratio = total / 35
        n_pdf = max(1, round(15 * ratio))
        n_jpg = max(1, round(10 * ratio))
        n_docx = max(1, round(5 * ratio))
        n_bad = max(1, round(5 * ratio))

    # Idempotent: clear and recreate output dir
    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    random.seed(42)

    generated = 0
    pdf_count = 0
    jpg_count = 0
    docx_count = 0
    bad_count = 0

    # Generate regular PDFs (indices 1..n_pdf-n_bad_pdf)
    n_bad_pdf = 3  # 3 bad cards in PDF format
    n_bad_docx = n_bad - n_bad_pdf  # rest as DOCX

    print(f'Generating {total} synthetic operational cards → {output_dir}')
    print()

    # Good PDFs
    for i in range(1, n_pdf - n_bad_pdf + 1):
        card = generate_card_data(i)
        fname = f'card_pdf_{i:02d}.pdf'
        fpath = output_dir / fname
        generate_pdf(card, str(fpath))
        pdf_count += 1
        generated += 1
        print(f'Generated {generated}/{total}: {fname}')

    # Bad PDFs
    for i in range(1, n_bad_pdf + 1):
        card = generate_card_data(100 + i, bad=True)
        fname = f'card_bad_pdf_{i:02d}.pdf'
        fpath = output_dir / fname
        generate_pdf(card, str(fpath))
        pdf_count += 1
        bad_count += 1
        generated += 1
        print(f'Generated {generated}/{total}: {fname}  [incomplete/bad]')

    # Scan-like JPGs
    for i in range(1, n_jpg + 1):
        card = generate_card_data(200 + i)
        fname = f'card_scan_{i:02d}.jpg'
        fpath = output_dir / fname
        generate_scan_jpg(card, str(fpath))
        jpg_count += 1
        generated += 1
        print(f'Generated {generated}/{total}: {fname}')

    # Good DOCXs
    for i in range(1, n_docx - n_bad_docx + 1):
        card = generate_card_data(300 + i)
        fname = f'card_docx_{i:02d}.docx'
        fpath = output_dir / fname
        generate_docx(card, str(fpath))
        docx_count += 1
        generated += 1
        print(f'Generated {generated}/{total}: {fname}')

    # Bad DOCXs
    for i in range(1, n_bad_docx + 1):
        card = generate_card_data(400 + i, bad=True)
        fname = f'card_bad_docx_{i:02d}.docx'
        fpath = output_dir / fname
        generate_docx(card, str(fpath))
        docx_count += 1
        bad_count += 1
        generated += 1
        print(f'Generated {generated}/{total}: {fname}  [incomplete/bad]')

    print()
    print('=' * 50)
    print('SUMMARY')
    print('=' * 50)
    print(f'Total generated : {generated}')
    print(f'  PDFs          : {pdf_count}  (including {n_bad_pdf} bad)')
    print(f'  Scan JPGs     : {jpg_count}')
    print(f'  DOCXs         : {docx_count}  (including {n_bad_docx} bad)')
    print(f'  Bad/incomplete: {bad_count}')
    print(f'Output dir      : {output_dir.resolve()}')

    # Verify file count
    actual_files = list(output_dir.iterdir())
    print(f'Files on disk   : {len(actual_files)}')
    if len(actual_files) != generated:
        print(f'WARNING: expected {generated} files but found {len(actual_files)}')


if __name__ == '__main__':
    main()
